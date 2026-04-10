#!/usr/bin/env python3
"""
Build a branded IOTTIVE DOCX from parsed document content.
Uses the Sample-SOW.docx as the base template — copies cover page (page 1) as-is,
then appends the parsed content starting from page 2 with IOTTIVE branding.
"""

import sys
import json
import copy
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_font(run, name="Lexend Light", size=None, bold=False, italic=False, color=None):
    """Apply IOTTIVE brand font to a run."""
    run.font.name = name
    # Set the east-asia and complex script fonts too
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), name)

    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_section_title(doc, text):
    """Add a right-aligned section title like in the SOW (e.g., 'Executive Summary')."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run(text)
    set_font(run, name="Lexend Medium", size=14)


def add_heading_h3(doc, text):
    """Add a Heading 3 style subheading."""
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.space_before = Pt(14)
    pf.space_after = Pt(4)
    pf.left_indent = Inches(0.5)
    run = para.add_run(text)
    set_font(run, name="Lexend Light", size=13, color=(0, 0, 0))
    run.font.bold = True


def add_paragraph(doc, runs_data, justify=True, space_before=True, space_after=True):
    """Add a body paragraph with IOTTIVE branding."""
    para = doc.add_paragraph()
    if justify:
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = para.paragraph_format
    if space_before:
        pf.space_before = Pt(12)
    if space_after:
        pf.space_after = Pt(12)

    for run_data in runs_data:
        text = run_data.get("text", "")
        if not text:
            continue
        run = para.add_run(text)
        is_bold = run_data.get("bold", False)
        is_italic = run_data.get("italic", False)
        set_font(run, name="Lexend Light", bold=is_bold, italic=is_italic)


def add_bullet(doc, runs_data):
    """Add a bullet point with IOTTIVE branding."""
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.space_before = Pt(12)
    pf.left_indent = Inches(0.5)

    # Add bullet character
    bullet_run = para.add_run("•  ")
    set_font(bullet_run, name="Lexend Light")

    for run_data in runs_data:
        text = run_data.get("text", "")
        if not text:
            continue
        run = para.add_run(text)
        is_bold = run_data.get("bold", False)
        is_italic = run_data.get("italic", False)
        set_font(run, name="Lexend Light", bold=is_bold, italic=is_italic)


def set_cell_border(cell, **kwargs):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, vals in kwargs.items():
        element = OxmlElement(f'w:{edge}')
        for attr, val in vals.items():
            element.set(qn(f'w:{attr}'), str(val))
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_table(doc, rows_data):
    """Add a table with IOTTIVE branding."""
    if not rows_data or len(rows_data) == 0:
        return

    num_cols = max(len(r.get("cells", [])) for r in rows_data)
    if num_cols == 0:
        return

    table = doc.add_table(rows=len(rows_data), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set table borders
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '6')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        borders.append(border)
    tblPr.append(borders)

    # Set table width to 100%
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)

    for ri, row_data in enumerate(rows_data):
        cells = row_data.get("cells", [])
        for ci in range(num_cols):
            cell = table.cell(ri, ci)
            cell_text = cells[ci] if ci < len(cells) else ""

            # Clear default paragraph
            for p in cell.paragraphs:
                p.clear()

            para = cell.paragraphs[0]
            run = para.add_run(cell_text)

            if ri == 0:
                # Header row: bold Lexend
                set_font(run, name="Lexend", bold=True)
            else:
                # Body row: Lexend Light
                set_font(run, name="Lexend Light")

    # Add spacing after table
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)


def add_page_break(doc):
    """Add a page break."""
    para = doc.add_paragraph()
    run = para.add_run()
    from docx.enum.text import WD_BREAK
    run.add_break(WD_BREAK.PAGE)


def build_branded_docx(data_path, template_path, output_path):
    """Main function: build branded DOCX."""
    with open(data_path, 'r') as f:
        data = json.load(f)

    elements = data.get("elements", [])
    client_name = data.get("clientName", "Client Name")
    contact_name = data.get("contactName", "Contact Name")
    contact_email = data.get("contactEmail", "email@example.com")
    project_title = data.get("projectTitle", "Technical & Commercial Proposal")

    # Load the brand template (which has the cover page, headers, footers, images)
    doc = Document(template_path)

    # Remove everything after the cover page (after paragraph index 5 which has the page break)
    # We'll keep paragraphs 0-5 (the cover page) and remove everything after
    body = doc.element.body
    children = list(body)

    # Find the first page break (cover page ends there)
    cover_end_idx = None
    para_count = 0
    for idx, child in enumerate(children):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            # Check for page break in this paragraph
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            breaks = child.findall('.//w:br', ns)
            for br in breaks:
                br_type = br.get(qn('w:type'))
                if br_type == 'page':
                    cover_end_idx = idx
                    break
            if cover_end_idx is not None:
                break
            para_count += 1

    # Remove everything after cover page
    if cover_end_idx is not None:
        elements_to_remove = children[cover_end_idx + 1:]
        # Keep the sectPr (last element, section properties)
        sect_pr = None
        for child in reversed(children):
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag == 'sectPr':
                sect_pr = child
                break

        for elem in elements_to_remove:
            tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
            if tag != 'sectPr':
                body.remove(elem)

    # Update cover page text: replace client name, contact, email, date
    import re
    from datetime import datetime

    today = datetime.now().strftime("%B %d, %Y")

    for para in doc.paragraphs:
        for run in para.runs:
            # Replace client name on cover
            if "RK Consultancy" in run.text:
                run.text = run.text.replace("RK Consultancy", client_name)
            if "Jay Solanki" in run.text:
                run.text = run.text.replace("Jay Solanki", contact_name)
            if "Jay@rkvisaconsultant.com" in run.text:
                run.text = run.text.replace("Jay@rkvisaconsultant.com", contact_email)
            if "December 30, 2025" in run.text:
                run.text = run.text.replace("December 30, 2025", today)

    # Update header text
    for section in doc.sections:
        header = section.header
        if header:
            for para in header.paragraphs:
                for run in para.runs:
                    if "RK Consultancy" in run.text:
                        run.text = run.text.replace("RK Consultancy", client_name)
                    if "CRM" in run.text:
                        run.text = run.text.replace("CRM", project_title.split()[0] if project_title else "")

    # Now add content from the parsed document starting from page 2
    # Group elements into sections for page breaks
    current_section_has_content = False
    is_first_section = True

    for i, elem in enumerate(elements):
        etype = elem.get("type")
        runs_data = elem.get("runs", [])
        text = elem.get("text", "")

        if not text and not runs_data and etype != "table":
            continue

        if etype == "heading":
            level = elem.get("level", 2)

            if level == 1:
                # Major section heading — right-aligned like SOW
                if not is_first_section and current_section_has_content:
                    add_page_break(doc)
                is_first_section = False
                current_section_has_content = False
                add_section_title(doc, text)
                doc.add_paragraph()  # spacing
            elif level == 2:
                add_section_title(doc, text)
                doc.add_paragraph()
            else:
                add_heading_h3(doc, text)

            current_section_has_content = True

        elif etype == "paragraph":
            if runs_data:
                add_paragraph(doc, runs_data)
            elif text:
                add_paragraph(doc, [{"text": text}])
            current_section_has_content = True

        elif etype == "bullet":
            if runs_data:
                add_bullet(doc, runs_data)
            elif text:
                add_bullet(doc, [{"text": text}])
            current_section_has_content = True

        elif etype == "table":
            rows = elem.get("rows", [])
            if rows:
                add_table(doc, rows)
            current_section_has_content = True

    doc.save(output_path)


if __name__ == "__main__":
    if len(sys.argv) != 4:
        print("Usage: build_docx.py <data.json> <template.docx> <output.docx>")
        sys.exit(1)

    build_branded_docx(sys.argv[1], sys.argv[2], sys.argv[3])
